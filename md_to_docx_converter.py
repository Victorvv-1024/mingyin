#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Document Converter
Converts the Chinese sales forecasting report from Markdown to Word format
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return heading

def add_paragraph_with_formatting(doc, text, style=None):
    """Add a paragraph with optional formatting"""
    # Remove markdown formatting
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
    text = re.sub(r'`(.*?)`', r'\1', text)        # Code
    
    para = doc.add_paragraph(text)
    if style:
        para.style = style
    return para

def create_table_from_markdown(doc, table_lines):
    """Create a Word table from markdown table lines"""
    if len(table_lines) < 3:  # Need at least header, separator, and one data row
        return
    
    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Parse data rows (skip separator line)
    data_rows = []
    for line in table_lines[2:]:
        if line.strip():
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(row) == len(headers):
                data_rows.append(row)
    
    if not data_rows:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        # Make header bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for row_data in data_rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            if i < len(row.cells):
                row.cells[i].text = cell_data
    
    return table

def convert_md_to_docx(md_file_path, docx_file_path):
    """Convert markdown file to Word document"""
    
    # Read the markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = content.split('\n')
    i = 0
    current_table_lines = []
    in_code_block = False
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                if current_table_lines:
                    code_text = '\n'.join(current_table_lines)
                    para = doc.add_paragraph(code_text)
                    para.style = 'Intense Quote'
                    current_table_lines = []
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            current_table_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line[1:]:
            current_table_lines.append(line)
            # Look ahead for more table lines
            j = i + 1
            while j < len(lines) and (lines[j].strip().startswith('|') or lines[j].strip().startswith('-')):
                current_table_lines.append(lines[j].strip())
                j += 1
            
            # Create table
            create_table_from_markdown(doc, current_table_lines)
            current_table_lines = []
            i = j
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('#').strip()
            add_heading(doc, title, min(level, 9))  # Word supports up to 9 heading levels
        
        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_page_break()
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            para = add_paragraph_with_formatting(doc, text)
            para.style = 'List Bullet'
        
        # Handle numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            para = add_paragraph_with_formatting(doc, text)
            para.style = 'List Number'
        
        # Handle bold/italic patterns
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = True
        
        # Handle regular paragraphs
        elif line and not line.startswith('#'):
            add_paragraph_with_formatting(doc, line)
        
        # Handle empty lines (add space)
        elif not line:
            doc.add_paragraph()
        
        i += 1
    
    # Save the document
    doc.save(docx_file_path)
    print(f"✅ Successfully converted {md_file_path} to {docx_file_path}")

if __name__ == "__main__":
    # Convert the Chinese report
    md_file = "销售预测模型结果分析报告_更新版.md"
    docx_file = "销售预测模型结果分析报告_更新版.docx"
    
    try:
        convert_md_to_docx(md_file, docx_file)
        print(f"📄 Word document created: {docx_file}")
        print("🎯 The document is ready for further formatting in Microsoft Word if needed")
    except Exception as e:
        print(f"❌ Error during conversion: {e}")
        import traceback
        traceback.print_exc() 